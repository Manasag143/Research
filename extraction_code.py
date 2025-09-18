import pandas as pd
import pdfplumber
import re
import os

def extract_contingent_liabilities_tables(pdf_path):
    """
    Simple function to extract contingent liabilities tables from PDF
    Using the correct physical page mappings:
    - Logical page 346 → Physical page 176
    - Logical page 423 → Physical page 214
    
    Args:
        pdf_path (str): Path to your PDF file
        
    Returns:
        dict: Simple results with tables and text
    """
    
    # Correct page mappings as provided
    pages_to_extract = {
        'consolidated_page': 176,  # Contains logical page 346
        'standalone_page': 214     # Contains logical page 423
    }
    
    logical_pages = {
        'consolidated_page': 346,
        'standalone_page': 423
    }
    
    print(f"Using correct page mappings:")
    print(f"Logical page 346 → Physical page 176")
    print(f"Logical page 423 → Physical page 214")
    
    # First, check PDF and get total pages
    try:
        with pdfplumber.open(pdf_path) as pdf:
            total_pages = len(pdf.pages)
            print(f"PDF opened successfully. Total physical pages: {total_pages}")
    except Exception as e:
        print(f"✗ Error opening PDF: {e}")
        return {}
    
    # Validate pages exist
    valid_pages = {}
    for section_name, page_num in pages_to_extract.items():
        if page_num <= total_pages:
            valid_pages[section_name] = page_num
            print(f"✓ Will extract {section_name.replace('_', ' ')} from physical page {page_num}")
        else:
            print(f"✗ Physical page {page_num} doesn't exist (PDF has {total_pages} pages)")
    
    if not valid_pages:
        print("✗ No valid pages found.")
        return {}
    
    results = {}
    
    print("Starting extraction...")
    
    for section_name, page_num in valid_pages.items():
        print(f"\nProcessing {section_name.replace('_', ' ')} (Physical page {page_num})...")
        
        section_results = {
            'logical_page': logical_pages[section_name],
            'physical_page': page_num,
            'text_found': '',
            'tables': []
        }
        
        try:
            # Step 1: Extract text from the page and separate left/right columns
            with pdfplumber.open(pdf_path) as pdf:
                page = pdf.pages[page_num - 1]  # Convert to 0-indexed
                
                # Get page dimensions
                page_width = page.width
                page_height = page.height
                half_width = page_width / 2
                
                # Extract text from left and right halves separately
                left_bbox = (0, 0, half_width, page_height)
                right_bbox = (half_width, 0, page_width, page_height)
                
                left_text = page.within_bbox(left_bbox).extract_text() or ""
                right_text = page.within_bbox(right_bbox).extract_text() or ""
                
                print(f"  Left column: {len(left_text)} characters")
                print(f"  Right column: {len(right_text)} characters")
                
                # Determine which section we're looking for
                if 'consolidated' in section_name:
                    target_main_heading = "notes to consolidated financial statement"
                    target_sub_heading = "contingent liabilit"
                    print(f"  Looking for: '{target_main_heading}' -> 'b) {target_sub_heading}'")
                else:
                    target_main_heading = "notes to standalone financial statement"
                    target_sub_heading = "contingent liabilit"
                    print(f"  Looking for: '{target_main_heading}' -> 'b) {target_sub_heading}'")
                
                # Check which column contains our target section
                relevant_text = ""
                found_section = False
                column_used = ""
                
                # Check left column first
                left_lower = left_text.lower()
                if (target_main_heading.lower() in left_lower and target_sub_heading in left_lower):
                    relevant_text = left_text
                    column_used = "LEFT"
                    print(f"  ✓ Found target section in LEFT column")
                # Check right column
                elif right_text:
                    right_lower = right_text.lower()
                    if (target_main_heading.lower() in right_lower and target_sub_heading in right_lower):
                        relevant_text = right_text
                        column_used = "RIGHT"
                        print(f"  ✓ Found target section in RIGHT column")
                
                if not relevant_text:
                    print(f"  ⚠ Target section not found in either column, checking for just contingent liabilities...")
                    # Fallback: look for just contingent liabilities in either column
                    if target_sub_heading in left_lower:
                        relevant_text = left_text
                        column_used = "LEFT"
                        print(f"  ⚠ Found only contingent liabilities in LEFT column")
                    elif target_sub_heading in right_lower:
                        relevant_text = right_text
                        column_used = "RIGHT"
                        print(f"  ⚠ Found only contingent liabilities in RIGHT column")
                
                if relevant_text:
                    found_section = True
                else:
                    print(f"  ✗ Target section not found in either column, using full page")
                    relevant_text = left_text + "\n\n" + right_text  # Combine both if nothing found
                    column_used = "BOTH"
                
                # Step 2: Find contingent liabilities section within the relevant column
                if found_section and relevant_text:
                    text_lower = relevant_text.lower()
                    
                    # Look for contingent liabilities heading patterns
                    patterns = [
                        r'b\)\s*contingent\s+liabilit[ieysn]*',
                        r'\(b\)\s*contingent\s+liabilit[ieysn]*',
                        r'contingent\s+liabilit[ieysn]*'
                    ]
                    
                    section_extracted = False
                    for pattern in patterns:
                        match = re.search(pattern, text_lower)
                        if match:
                            section_extracted = True
                            
                            # Start from the beginning of the heading
                            start = match.start()
                            
                            # Find the end of this section
                            end_patterns = [
                                r'\bc\)\s*',  # Next sub-section c)
                                r'\(c\)\s*',  # Next sub-section (c)
                                r'\bd\)\s*',  # Next sub-section d)
                                r'\(d\)\s*',  # Next sub-section (d)
                                r'\n\s*\d+\.\s*',  # Next numbered section
                                r'\n\s*[A-Z][A-Z\s]{10,}\n'  # Next major heading
                            ]
                            
                            # Look for section end starting from after the heading
                            search_start = match.end() + 50
                            end = len(relevant_text)  # Default to end of column
                            
                            for end_pattern in end_patterns:
                                end_match = re.search(end_pattern, text_lower[search_start:])
                                if end_match:
                                    end = search_start + end_match.start()
                                    break
                            
                            section_results['text_found'] = relevant_text[start:end].strip()
                            print(f"  ✓ Extracted contingent liabilities section from {column_used} column")
                            print(f"  ✓ Section length: {len(section_results['text_found'])} characters")
                            break
                    
                    if not section_extracted:
                        # If specific section not found, use the relevant column
                        section_results['text_found'] = relevant_text
                        print(f"  ⚠ Used entire {column_used} column text")
                
                else:
                    print(f"  ⚠ Using combined text from both columns")
                    section_results['text_found'] = relevant_text
            
                # Step 3: Extract tables from the correct column only
                try:
                    print(f"  Extracting tables from {column_used} column...")
                    
                    # Focus table extraction on the specific column that contains our section
                    with pdfplumber.open(pdf_path) as pdf:
                        page = pdf.pages[page_num - 1]
                        
                        # Define the area to extract tables from
                        if column_used == "LEFT":
                            table_area = (0, 0, half_width, page_height)
                            cropped_page = page.within_bbox(table_area)
                        elif column_used == "RIGHT":
                            table_area = (half_width, 0, page_width, page_height)
                            cropped_page = page.within_bbox(table_area)
                        else:  # BOTH
                            cropped_page = page
                        
                        # Try pdfplumber table extraction first
                        page_tables = cropped_page.extract_tables()
                        
                        if page_tables:
                            print(f"  ✓ pdfplumber found {len(page_tables)} structured tables in {column_used} column")
                            for i, table_data in enumerate(page_tables):
                                if table_data and len(table_data) > 1:
                                    df = pd.DataFrame(table_data[1:], columns=table_data[0])
                                    df = df.dropna(how='all').dropna(axis=1, how='all').fillna('')
                                    
                                    if not df.empty:
                                        section_results['tables'].append(df)
                                        print(f"  ✓ Added table {i+1}: {df.shape}")
                        else:
                            print(f"  ⚠ No structured tables found in {column_used} column")
                        
                        # Text-based table extraction from the specific section
                        print(f"  Analyzing text from {column_used} column for tabular data...")
                        text_to_analyze = section_results['text_found']
                        lines = text_to_analyze.split('\n')
                        potential_table_lines = []
                        
                        for line in lines:
                            line = line.strip()
                            # Look for lines that are likely table rows
                            if (10 < len(line) < 200 and 
                                (line.count('₹') >= 1 or 
                                 line.count('Rs.') >= 1 or
                                 'crore' in line.lower() or
                                 'lakh' in line.lower() or
                                 len([word for word in line.split() if word.replace(',', '').replace('.', '').replace('(', '').replace(')', '').isdigit()]) >= 1)):
                                
                                # Skip header-like lines
                                line_lower = line.lower()
                                if not any(skip_word in line_lower for skip_word in 
                                         ['note', 'contingent', 'liabilit', 'statement', 'financial', 'as on', 'previous year']):
                                    potential_table_lines.append(line)
                        
                        if potential_table_lines:
                            print(f"  ✓ Found {len(potential_table_lines)} potential financial data lines")
                            
                            # Create table from text lines
                            table_data = []
                            for line in potential_table_lines:
                                # Split by significant spacing
                                row = re.split(r'\s{2,}', line.strip())
                                
                                # Clean up the row
                                cleaned_row = [cell.strip() for cell in row if cell.strip()]
                                
                                if len(cleaned_row) >= 2:  # At least description + amount
                                    table_data.append(cleaned_row)
                            
                            if table_data:
                                # Determine maximum number of columns
                                max_cols = max(len(row) for row in table_data)
                                
                                # Pad shorter rows
                                for row in table_data:
                                    while len(row) < max_cols:
                                        row.append('')
                                
                                # Create DataFrame with intelligent column names
                                if max_cols == 2:
                                    col_names = ['Description', 'Amount']
                                elif max_cols == 3:
                                    col_names = ['Description', 'Current Year', 'Previous Year']
                                elif max_cols == 4:
                                    col_names = ['Description', 'Current Year', 'Previous Year', 'Notes']
                                else:
                                    col_names = [f'Column_{i+1}' for i in range(max_cols)]
                                
                                df = pd.DataFrame(table_data, columns=col_names)
                                df = df.replace('', pd.NA).dropna(how='all').fillna('')
                                
                                if not df.empty:
                                    section_results['tables'].append(df)
                                    print(f"  ✓ Created table from text: {df.shape[0]} rows × {df.shape[1]} columns")
                                    print(f"  ✓ Column structure: {list(df.columns)}")
                        
                        if not section_results['tables']:
                            print(f"  ℹ No tables extracted from {column_used} column")
                
                except Exception as e:
                    print(f"  ✗ Table extraction failed: {e}")
                    print(f"  ℹ Continuing with text extraction only...")
            
            results[section_name] = section_results
            print(f"  Done: Found {len(section_results['tables'])} tables")
            
        except Exception as e:
            print(f"  ✗ Failed to process page {page_num}: {e}")
            results[section_name] = section_results
    
    return results

def save_results_to_word(results, output_file="contingent_liabilities_extracted.docx"):
    """
    Save clean extracted text as tables in Word document
    
    Args:
        results (dict): Extraction results
        output_file (str): Output Word file name
    """
    
    try:
        from docx import Document
        print(f"\nSaving results to {output_file}...")
        
        # Create a new Word document
        doc = Document()
        
        # Add title
        doc.add_heading('Contingent Liabilities Data', 0)
        
        # Process each section
        for section_name, data in results.items():
            section_title = section_name.replace('_', ' ').title()
            
            print(f"Processing {section_title} for Word document...")
            
            # Add section heading
            doc.add_heading(f'{section_title}', level=1)
            doc.add_paragraph(f"Page {data['physical_page']} (Logical page {data['logical_page']})")
            
            # Get the extracted text
            text = data['text_found']
            
            if not text:
                doc.add_paragraph("No text was extracted for this section.")
                print(f"  ⚠ No text found for {section_title}")
                continue
            
            print(f"  Text length: {len(text)} characters")
            
            # First, add the raw extracted text for debugging
            doc.add_heading('Extracted Text:', level=2)
            doc.add_paragraph(text)
            
            # Now try to create table from the text
            doc.add_heading('Processed Table:', level=2)
            
            lines = text.split('\n')
            print(f"  Total lines: {len(lines)}")
            
            # Less restrictive filtering - keep more lines
            table_lines = []
            
            for i, line in enumerate(lines):
                line = line.strip()
                
                # Skip completely empty lines
                if not line:
                    continue
                
                print(f"  Line {i+1}: {line[:80]}...")  # Debug print
                
                # Less restrictive - keep lines that might have financial data
                # Look for any line with numbers, currency, or financial terms
                has_financial_data = (
                    any(char.isdigit() for char in line) or  # Has numbers
                    '₹' in line or 'Rs.' in line or 'rs.' in line.lower() or  # Currency
                    'crore' in line.lower() or 'lakh' in line.lower() or  # Amount terms
                    'million' in line.lower() or  # Amount terms
                    len(line) > 10  # Not too short
                )
                
                # Skip obvious headers but be less restrictive
                is_header = any(header in line.lower() for header in 
                              ['contingent liabilities', 'notes to', 'financial statement'])
                
                if has_financial_data and not is_header:
                    table_lines.append(line)
                    print(f"    ✓ Keeping line: {line}")
                
                # Stop at total but continue to show what we found
                if 'total' in line.lower() and has_financial_data:
                    print(f"    ✓ Found total line: {line}")
                    break
            
            print(f"  ✓ Found {len(table_lines)} potential data lines")
            
            # Create table if we have data
            if table_lines:
                # Simple table creation
                table_data = []
                
                for line in table_lines:
                    # Try different splitting methods
                    columns = re.split(r'\s{2,}', line)  # Split by 2+ spaces
                    
                    # Clean columns
                    clean_columns = [col.strip() for col in columns if col.strip()]
                    
                    if clean_columns:  # If we have any columns
                        # Pad to 3 columns
                        while len(clean_columns) < 3:
                            clean_columns.append('')
                        
                        table_data.append(clean_columns[:3])  # Take first 3
                
                if table_data:
                    # Create Word table
                    word_table = doc.add_table(rows=1, cols=3)
                    word_table.style = 'Table Grid'
                    
                    # Headers
                    headers = ['Description', 'Amount 1', 'Amount 2']
                    header_cells = word_table.rows[0].cells
                    for i, header in enumerate(headers):
                        header_cells[i].text = header
                        for paragraph in header_cells[i].paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    
                    # Add data
                    for row_data in table_data:
                        row_cells = word_table.add_row().cells
                        for i, value in enumerate(row_data):
                            row_cells[i].text = str(value)
                    
                    doc.add_paragraph()
                    print(f"  ✓ Created table with {len(table_data)} rows")
                
                else:
                    doc.add_paragraph("Could not parse the data into table format.")
                    print(f"  ⚠ Could not create table structure")
            
            else:
                doc.add_paragraph("No financial data lines found in the extracted text.")
                print(f"  ⚠ No financial data lines found")
            
            # Add page break between sections
            if section_name != list(results.keys())[-1]:
                doc.add_page_break()
        
        # Save
        doc.save(output_file)
        print(f"✓ Word document saved: {output_file}")
        return output_file
        
    except ImportError:
        print("✗ python-docx not installed. Install with: pip install python-docx")
        return None
    except Exception as e:
        print(f"✗ Error saving Word document: {e}")
        return None

def search_for_section(pdf_path, section_type, total_pages):
    """
    Search for contingent liabilities sections in the PDF
    
    Args:
        pdf_path (str): Path to PDF
        section_type (str): 'consolidated' or 'standalone'
        total_pages (int): Total pages in PDF
        
    Returns:
        int: Page number if found, None otherwise
    """
    
    search_terms = {
        'consolidated': ['consolidated financial statement', 'consolidated financial statements'],
        'standalone': ['standalone financial statement', 'standalone financial statements']
    }
    
    print(f"  Searching for {section_type} section...")
    
    try:
        with pdfplumber.open(pdf_path) as pdf:
            # Search in reasonable range (avoid searching entire PDF for large files)
            start_page = max(1, total_pages - 100)  # Search last 100 pages
            end_page = total_pages
            
            for page_num in range(start_page, end_page + 1):
                try:
                    page = pdf.pages[page_num - 1]
                    text = page.extract_text()
                    if text:
                        text_lower = text.lower()
                        
                        # Check for section type and contingent liabilities
                        section_found = any(term in text_lower for term in search_terms[section_type])
                        contingent_found = 'contingent liabilit' in text_lower
                        
                        if section_found and contingent_found:
                            print(f"  ✓ Found {section_type} section on page {page_num}")
                            return page_num
                            
                except Exception as e:
                    continue  # Skip problematic pages
                    
    except Exception as e:
        print(f"  ✗ Search failed: {e}")
    
    print(f"  ✗ {section_type} section not found")
    return None

def main():
    """
    Simple main function to run everything
    """
    
    # Set your PDF path here - CHANGE THIS TO YOUR ACTUAL PDF PATH
    pdf_path = input("Enter the path to your PDF file: ").strip().strip('"').strip("'")
    
    # Alternative: uncomment and set your PDF path directly
    # pdf_path = r"C:\path\to\your\pdf\file.pdf"  # Windows
    # pdf_path = "/path/to/your/pdf/file.pdf"     # Mac/Linux
    
    if not pdf_path or not os.path.exists(pdf_path):
        print("❌ PDF file not found. Please check the path.")
        print("Example paths:")
        print("  Windows: C:\\Documents\\your_file.pdf")
        print("  Mac/Linux: /Users/username/Documents/your_file.pdf")
        return
    
    print("PDF Contingent Liabilities Extractor")
    print("=" * 40)
    
    # Extract data using the correct physical page numbers
    results = extract_contingent_liabilities_tables(pdf_path)
    
    if not results:
        print("No sections found. Let me help you find the right pages...")
        
        # Show some sample pages to help user identify correct pages
        try:
            with pdfplumber.open(pdf_path) as pdf:
                total_pages = len(pdf.pages)
                print(f"\nYour PDF has {total_pages} pages.")
                print("Let me show you some pages that mention 'contingent liabilities':\n")
                
                found_pages = []
                for page_num in range(1, min(total_pages + 1, 50)):  # Check first 50 pages
                    try:
                        page = pdf.pages[page_num - 1]
                        text = page.extract_text()
                        if text and 'contingent liabilit' in text.lower():
                            found_pages.append(page_num)
                            print(f"Page {page_num}: Found 'contingent liabilities'")
                            # Show first 200 chars to help identify
                            start_pos = text.lower().find('contingent liabilit')
                            if start_pos >= 0:
                                snippet = text[max(0, start_pos-50):start_pos+150]
                                print(f"  Context: ...{snippet}...")
                            print()
                            
                            if len(found_pages) >= 5:  # Limit output
                                break
                    except:
                        continue
                
                if found_pages:
                    print(f"Found 'contingent liabilities' on pages: {found_pages}")
                    print("\nTo extract specific pages, modify the code like this:")
                    print("results = extract_contingent_liabilities_tables(pdf_path,")
                    print(f"                                               consolidated_page={found_pages[0]},")
                    if len(found_pages) > 1:
                        print(f"                                               standalone_page={found_pages[1]})")
                    else:
                        print("                                               standalone_page=None)")
                else:
                    print("No pages found with 'contingent liabilities'. Please check your PDF.")
                    
        except Exception as e:
            print(f"Error analyzing PDF: {e}")
        
        return
    
    # Save to Word document
    if results:
        save_results_to_word(results)
    
    # Print simple summary
    print("\n" + "=" * 40)
    print("EXTRACTION COMPLETE")
    print("=" * 40)
    
    for section_name, data in results.items():
        section_title = section_name.replace('_', ' ').title()
        print(f"{section_title}:")
        print(f"  Logical page: {data['logical_page']}")
        print(f"  Physical page: {data['physical_page']}")
        print(f"  Tables found: {len(data['tables'])}")
        print(f"  Text extracted: {len(data['text_found'])} characters")
        
        if data['tables']:
            print("  Table shapes:")
            for i, table in enumerate(data['tables']):
                print(f"    Table {i+1}: {table.shape[0]} rows × {table.shape[1]} columns")
        print()
    
    print("✓ All done!")

# Run the extraction
if __name__ == "__main__":
    main()
